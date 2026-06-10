"""Build an editable Word document sales flyer for Mobile Numerology course.

Output: /app/frontend/public/mobile-numerology-flyer.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ---------- colour palette ----------
RED      = (0xC0, 0x00, 0x20)
DARK_RED = (0x5B, 0x0B, 0x1F)
GOLD     = (0xB8, 0x88, 0x1A)
DARK     = (0x2A, 0x1A, 0x2C)
WHITE    = (0xFF, 0xFF, 0xFF)
CREAM    = "FFF8EC"

BONUSES = [
    ("Mobile Numerology Software",  "Lifetime access — calculate any number instantly", "₹9,999"),
    ("Recording — Lifetime Access", "Re-watch every live session anytime, forever",      "₹7,999"),
    ("Tried & Tested Vaastu Remedies",      "For Money · Health · Relationships",        "₹4,999"),
    ("Tried & Tested Numerology Remedies",  "Personal mantras & rituals from 14+ years", "₹4,999"),
    ("Switch Words Pack",                   "Hand-picked words for abundance & calm",    "₹2,999"),
    ("Name Numerology Calculator",          "Premium calculator — unlimited usage",      "₹2,999"),
    ("Personal Year Calculator",            "Year-by-year forecast tool with insights",  "₹2,999"),
]

PHONE_DISPLAY = "+91 99290 59153"


# ---------- helpers ----------
def add_run(p, text, *, size=11, bold=False, italic=False, color=None, font="Calibri"):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, *, color="D4AF37", size_eighth_pt=12, style="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), style)
        b.set(qn("w:sz"), str(size_eighth_pt))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_horizontal_line(doc, color="D4AF37", size_eighth_pt=8):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighth_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(2)


# ---------- build the document ----------
doc = Document()

# Page margins (compact A4)
section = doc.sections[0]
section.top_margin    = Cm(1.2)
section.bottom_margin = Cm(1.0)
section.left_margin   = Cm(1.6)
section.right_margin  = Cm(1.6)

# Default style — readable on print
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ====== HEADER ======
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "NEWALKKAR  SAANDIIP", size=22, bold=True, color=DARK_RED, font="Georgia")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "NUMEROLOGIST  ·  VAASTU CONSULTANT  ·  LIFE COACH", size=9, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(4)

add_horizontal_line(doc)

# ====== COURSE TITLE ======
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "FLAGSHIP COURSE", size=8, bold=True, color=WHITE, font="Calibri")
# wrap label in red highlight via cell? Just use color contrast — set a tiny single-cell table:
# Simpler: leave as bold caption above title
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Mobile Numerology Complete Course", size=24, bold=True, color=DARK_RED, font="Georgia")
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p,
    "Decode the hidden vibration in your mobile number — and learn how a single number change "
    "can unlock ", size=11)
add_run(p, "Money · Health · Relationships", size=11, bold=True, color=DARK_RED)
add_run(p, ".", size=11)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "✓ 14+ Years Practice     ✓ 100,000+ Analyses     ✓ Lifetime Access",
        size=10, italic=True, color=GOLD, bold=True)
p.paragraph_format.space_after = Pt(8)

# ====== RED BONUS BANNER ======
banner = doc.add_table(rows=1, cols=1)
banner_cell = banner.rows[0].cells[0]
set_cell_bg(banner_cell, "C00020")
set_cell_borders(banner_cell, color="C00020", size_eighth_pt=4)
banner_p = banner_cell.paragraphs[0]
banner_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(banner_p, "🎁  ENROLL TODAY & UNLOCK  ", size=11, bold=True, color=WHITE)
add_run(banner_p, "FREE BONUSES WORTH ₹36,993+", size=18, bold=True, color=WHITE)
banner_p.paragraph_format.space_before = Pt(4)
banner_p.paragraph_format.space_after = Pt(4)

# spacer
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ====== BONUS LIST ======
p = doc.add_paragraph()
add_run(p, "YOUR FREE BONUSES (Included when you enroll today):",
        size=12, bold=True, color=DARK_RED)
p.paragraph_format.space_after = Pt(4)

for idx, (title, sub, price) in enumerate(BONUSES, start=1):
    # main bonus line
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"  {idx}.  ", size=11, bold=True, color=GOLD)
    add_run(p, "FREE  ", size=11, bold=True, color=RED)
    add_run(p, title, size=12, bold=True, color=DARK_RED)
    # sub-line
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, sub + "   ", size=10, color=DARK)
    add_run(p, "—  Value: ", size=10, italic=True, color=GOLD)
    add_run(p, price, size=10, bold=True, color=RED)

# Total value summary
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(2)
add_run(p, "TOTAL BONUS VALUE: ", size=11, bold=True, color=DARK_RED)
add_run(p, "₹36,993+", size=14, bold=True, color=RED)
add_run(p, "    —    YOUR PRICE TODAY: ₹_______", size=11, bold=True, color=DARK_RED)
p.paragraph_format.space_after = Pt(8)

# ====== QR + PHONE BLOCK ======
contact_table = doc.add_table(rows=1, cols=2)
contact_table.autofit = False

qr_cell, info_cell = contact_table.rows[0].cells
qr_cell.width = Cm(5.5)
info_cell.width = Cm(11.5)

# --- QR placeholder cell ---
set_cell_bg(qr_cell, CREAM)
set_cell_borders(qr_cell, color="D4AF37", size_eighth_pt=16, style="dashed")
qr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p = qr_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "\n", size=10)
add_run(p, "▢ ▢ ▢ ▢\n▢      ▢\n▢      ▢\n▢ ▢ ▢ ▢", size=12, bold=True, color=GOLD)
p = qr_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "\nPASTE YOUR\nPAYMENT  QR\nHERE", size=10, bold=True, color=DARK_RED)
p = qr_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "(right-click → Insert Picture)", size=8, italic=True, color=GOLD)

# --- Info cell ---
set_cell_bg(info_cell, "0F0518")
set_cell_borders(info_cell, color="D4AF37", size_eighth_pt=12)
info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p = info_cell.paragraphs[0]
add_run(p, "  RESERVE YOUR SEAT NOW", size=9, bold=True, color=(0xF3, 0xD0, 0x60))
p = info_cell.add_paragraph()
add_run(p, f"  📞  {PHONE_DISPLAY}", size=22, bold=True, color=WHITE, font="Georgia")
p = info_cell.add_paragraph()
add_run(p, "  Call or WhatsApp Newalkkar Saandiip ji directly — quote ",
        size=10, color=(0xE8, 0xDC, 0xC4))
add_run(p, "\"Mobile Numerology Course\"",
        size=10, bold=True, italic=True, color=(0xF3, 0xD0, 0x60))
add_run(p, " to lock-in today's launch price + all bonuses listed above.",
        size=10, color=(0xE8, 0xDC, 0xC4))
p = info_cell.add_paragraph()
add_run(p, "  ⏳  ONLY 200 SEATS  —  LIMITED-TIME LAUNCH OFFER",
        size=10, bold=True, color=(0xFF, 0x6B, 0x6B))

# spacer below contact block
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ====== FOOTER ======
add_horizontal_line(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "100,000+ Analyses   ·   14+ Years of Practice   ·   newalkkarsaandiip.in",
        size=9, bold=True, color=GOLD)

# ====== SAVE ======
import os
out_dir = "/app/frontend/public"
os.makedirs(out_dir, exist_ok=True)
out = os.path.join(out_dir, "mobile-numerology-flyer.docx")
doc.save(out)
print(f"Saved -> {out}")
